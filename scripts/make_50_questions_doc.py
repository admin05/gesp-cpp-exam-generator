from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


SRC = Path("/Users/cen/Downloads/20道题01.docx")
OUT = Path("/Users/cen/Documents/Codex/2026-05-18/4-5-1-2-3-4/20道题01_50题含答案.docx")


original_starts = [0, 10, 18, 32, 43, 54, 67, 81, 93, 110, 126, 138, 150, 160, 172, 178, 195, 204, 220, 232]

answers = {
    1: "B", 2: "B", 3: "A", 4: "C", 5: "B",
    6: "C", 7: "C", 8: "B", 9: "B", 10: "A",
    11: "C", 12: "B", 13: "B", 14: "B", 15: "C",
    16: "B", 17: "B", 18: "B", 19: "B", 20: "B",
}

new_questions = [
    (21, [
        "下面程序片段的输出结果是？",
        "int x = 0;",
        "for (int i = 1; i <= 6; i++) {",
        "    if (i % 2 == 0) x += i;",
        "    else x -= i;",
        "}",
        "cout << x;",
        "A. -3", "B. 3", "C. 6", "D. 12",
    ], "B"),
    (22, [
        "执行下面程序片段后，输出结果是？",
        "int a = 17, b = 5;",
        "cout << a / b + a % b;",
        "A. 3", "B. 5", "C. 6", "D. 8",
    ], "B"),
    (23, [
        "下面程序片段的输出结果是？",
        "int a = 4, b = 9, c = 6;",
        "int m = (a > b ? a : b);",
        "m = (m < c ? c : m);",
        "cout << m;",
        "A. 4", "B. 6", "C. 9", "D. 15",
    ], "C"),
    (24, [
        "下面程序片段的输出结果是？",
        "int cnt = 0;",
        "for (int i = 1; i <= 5; i++) {",
        "    for (int j = 1; j <= 5; j++) {",
        "        if (i + j == 6) cnt++;",
        "    }",
        "}",
        "cout << cnt;",
        "A. 4", "B. 5", "C. 6", "D. 10",
    ], "B"),
    (25, [
        "执行下面程序片段后，输出结果是？",
        "int s = 0;",
        "for (int i = 1; i <= 4; i++) {",
        "    for (int j = 1; j <= i; j++) {",
        "        s += i - j;",
        "    }",
        "}",
        "cout << s;",
        "A. 4", "B. 6", "C. 10", "D. 20",
    ], "C"),
    (26, [
        "已知数组：",
        "int a[7] = {2, 5, 1, 8, 4, 8, 3};",
        "int k = 0;",
        "for (int i = 1; i < 7; i++) {",
        "    if (a[i] >= a[k]) k = i;",
        "}",
        "cout << k;",
        "输出结果是？",
        "A. 3", "B. 4", "C. 5", "D. 6",
    ], "C"),
    (27, [
        "下面程序片段的输出结果是？",
        "int a[6] = {1, 3, 5, 7, 9, 11};",
        "int l = 0, r = 5, x = 7, ans = -1;",
        "while (l <= r) {",
        "    int mid = (l + r) / 2;",
        "    if (a[mid] == x) { ans = mid; break; }",
        "    else if (a[mid] < x) l = mid + 1;",
        "    else r = mid - 1;",
        "}",
        "cout << ans;",
        "A. 2", "B. 3", "C. 4", "D. -1",
    ], "B"),
    (28, [
        "下面程序片段执行后，数组 a 的前 5 个元素为？",
        "int a[5] = {5, 1, 4, 2, 3};",
        "for (int i = 1; i < 5; i++) {",
        "    int x = a[i], j = i - 1;",
        "    while (j >= 0 && a[j] > x) {",
        "        a[j + 1] = a[j];",
        "        j--;",
        "    }",
        "    a[j + 1] = x;",
        "}",
        "A. {5, 4, 3, 2, 1}", "B. {1, 2, 3, 4, 5}", "C. {1, 3, 2, 4, 5}", "D. {5, 1, 2, 3, 4}",
    ], "B"),
    (29, [
        "下面程序片段的输出结果是？",
        "int a[4] = {2, 4, 6, 8};",
        "int b[4];",
        "for (int i = 0; i < 4; i++) {",
        "    b[i] = a[(i + 1) % 4];",
        "}",
        "cout << b[0] << b[3];",
        "A. 24", "B. 42", "C. 82", "D. 48",
    ], "B"),
    (30, [
        "下面程序片段的输出结果是？",
        "int a[3][3] = {{2, 1, 3}, {0, 4, 5}, {7, 6, 8}};",
        "int s = 0;",
        "for (int i = 0; i < 3; i++) {",
        "    for (int j = 0; j < 3; j++) {",
        "        if (i < j) s += a[i][j];",
        "    }",
        "}",
        "cout << s;",
        "A. 9", "B. 12", "C. 15", "D. 18",
    ], "A"),
    (31, [
        "执行下面程序片段后，输出结果是？",
        "int a[3][4] = {{1,2,3,4},{5,6,7,8},{9,10,11,12}};",
        "int sum = 0;",
        "for (int i = 0; i < 3; i++) sum += a[i][i + 1];",
        "cout << sum;",
        "A. 18", "B. 21", "C. 24", "D. 27",
    ], "B"),
    (32, [
        "下面程序片段的输出结果是？",
        "string s = \"abcabc\";",
        "int cnt = 0;",
        "for (int i = 0; i + 1 < s.size(); i++) {",
        "    if (s[i] < s[i + 1]) cnt++;",
        "}",
        "cout << cnt;",
        "A. 2", "B. 3", "C. 4", "D. 5",
    ], "C"),
    (33, [
        "执行下面程序片段后，输出结果是？",
        "string s = \"10203\";",
        "int x = 0;",
        "for (int i = 0; i < s.size(); i++) {",
        "    if (s[i] != '0') x = x * 10 + (s[i] - '0');",
        "}",
        "cout << x;",
        "A. 123", "B. 10203", "C. 1203", "D. 6",
    ], "A"),
    (34, [
        "下面函数调用 g(10) 的返回值是？",
        "int g(int n) {",
        "    if (n == 0) return 0;",
        "    return n % 2 + g(n / 2);",
        "}",
        "A. 1", "B. 2", "C. 3", "D. 4",
    ], "B"),
    (35, [
        "下面函数调用 f(4) 的返回值是？",
        "int f(int n) {",
        "    if (n == 1) return 2;",
        "    return f(n - 1) * 2 + 1;",
        "}",
        "A. 7", "B. 15", "C. 23", "D. 31",
    ], "C"),
    (36, [
        "下面程序片段的输出结果是？",
        "int n = 30, cnt = 0;",
        "for (int i = 1; i * i <= n; i++) {",
        "    if (n % i == 0) {",
        "        cnt++;",
        "        if (i * i != n) cnt++;",
        "    }",
        "}",
        "cout << cnt;",
        "A. 6", "B. 8", "C. 10", "D. 12",
    ], "B"),
    (37, [
        "用筛法标记 2 到 30 的合数，最后未被标记的数有多少个？",
        "A. 8", "B. 9", "C. 10", "D. 11",
    ], "C"),
    (38, [
        "下面程序片段统计的是？",
        "int cnt = 0;",
        "for (int x = 1; x <= 20; x++) {",
        "    if (x % 3 == 0 && x % 5 != 0) cnt++;",
        "}",
        "cout << cnt;",
        "A. 4", "B. 5", "C. 6", "D. 7",
    ], "B"),
    (39, [
        "若用贪心策略解决“用面值 1、5、10、20 的纸币凑出 36 元且张数尽量少”，最少需要几张？",
        "A. 3", "B. 4", "C. 5", "D. 6",
    ], "B"),
    (40, [
        "有物品重量分别为 2、3、4、5、9，背包容量为 10。若只要求装入物品数量尽量多，应该选择的最大件数是？",
        "A. 2", "B. 3", "C. 4", "D. 5",
    ], "B"),
    (41, [
        "下面程序片段的输出结果是？",
        "int cnt = 0;",
        "for (int a = 0; a <= 2; a++) {",
        "    for (int b = 0; b <= 2; b++) {",
        "        if (a + b == 2) cnt++;",
        "    }",
        "}",
        "cout << cnt;",
        "A. 2", "B. 3", "C. 4", "D. 5",
    ], "B"),
    (42, [
        "从数字 1、2、3、4 中选出 3 个互不相同的数字组成三位数，其中偶数有多少个？",
        "A. 8", "B. 10", "C. 12", "D. 24",
    ], "C"),
    (43, [
        "下面递推程序的输出结果是？",
        "int a[6];",
        "a[1] = 1; a[2] = 2;",
        "for (int i = 3; i <= 5; i++) {",
        "    a[i] = a[i - 1] + 2 * a[i - 2];",
        "}",
        "cout << a[5];",
        "A. 10", "B. 16", "C. 22", "D. 26",
    ], "C"),
    (44, [
        "下面程序片段的输出结果是？",
        "int n = 12345, s = 0;",
        "while (n > 0) {",
        "    int d = n % 10;",
        "    if (d % 2 == 1) s = s * 10 + d;",
        "    n /= 10;",
        "}",
        "cout << s;",
        "A. 135", "B. 531", "C. 54321", "D. 246",
    ], "B"),
    (45, [
        "下面程序片段的输出结果是？",
        "int a = 0, b = 1;",
        "for (int i = 1; i <= 5; i++) {",
        "    int c = a + b;",
        "    a = b;",
        "    b = c;",
        "}",
        "cout << a;",
        "A. 3", "B. 5", "C. 8", "D. 13",
    ], "B"),
    (46, [
        "下面程序片段执行后，输出结果是？",
        "int x = 6;",
        "bool ok = (x % 2 == 0 && x % 3 == 0) || x % 5 == 0;",
        "cout << ok;",
        "A. 0", "B. 1", "C. true", "D. false",
    ], "B"),
    (47, [
        "甲、乙、丙三人参加比赛，三人名次互不相同。甲说：“我不是第一。”乙说：“甲是第三。”丙说：“乙不是第二。”三句话中恰有两句为真。谁是第一？",
        "A. 甲", "B. 乙", "C. 丙", "D. 无法确定",
    ], "C"),
    (48, [
        "一个袋子中有 2 个红球、3 个蓝球、5 个白球。随机取 1 个球，取到的不是白球的概率是？",
        "A. 1/5", "B. 1/2", "C. 3/5", "D. 4/5",
    ], "B"),
    (49, [
        "下面程序片段的输出结果是？",
        "int cnt = 0;",
        "for (int i = 1; i <= 4; i++) {",
        "    for (int j = 1; j <= 4; j++) {",
        "        if (i < j && (i + j) % 2 == 1) cnt++;",
        "    }",
        "}",
        "cout << cnt;",
        "A. 2", "B. 3", "C. 4", "D. 6",
    ], "C"),
    (50, [
        "下面程序片段的输出结果是？",
        "int a[5] = {2, 1, 3, 1, 2};",
        "int ans = 0;",
        "for (int i = 0; i < 5; i++) {",
        "    for (int j = i + 1; j < 5; j++) {",
        "        if (a[i] == a[j]) ans++;",
        "    }",
        "}",
        "cout << ans;",
        "A. 1", "B. 2", "C. 3", "D. 4",
    ], "B"),
]


def set_normal_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)


def add_lines(doc: Document, lines: list[str]) -> None:
    for line in lines:
        doc.add_paragraph(line)


def main() -> None:
    doc = Document(SRC)
    set_normal_font(doc)

    for n, idx in enumerate(original_starts, 1):
        p = doc.paragraphs[idx]
        text = p.text.strip()
        if text and not text.startswith(f"{n}."):
            p.text = f"{n}. {text}"

    for n, lines, answer in new_questions:
        doc.add_paragraph("")
        first, rest = lines[0], lines[1:]
        doc.add_paragraph(f"{n}. {first}")
        add_lines(doc, rest)
        answers[n] = answer

    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)
    title = doc.add_paragraph("答案")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)

    table = doc.add_table(rows=11, cols=5)
    for col in range(5):
        table.cell(0, col).text = f"题号/答案"
    for col in range(5):
        for row in range(1, 11):
            qn = col * 10 + row
            table.cell(row, col).text = f"{qn}. {answers[qn]}"

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
