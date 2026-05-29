from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
import re


OUT = "/Users/cen/Documents/Codex/2026-05-18/4-5-1-2-3-4/错题薄弱点专项训练50题_C++高亮缩进.docx"


questions = [
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int a = 8, b = 3;\na = a + b * 2;\nb = a / b + a % b;\ncout << a << \" \" << b;",
        "opts": ["A. 14 4", "B. 14 6", "C. 11 5", "D. 16 5"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int x = 4;\ncout << (x % 2 == 0 ? x * x : x + 10);",
        "opts": ["A. 4", "B. 8", "C. 14", "D. 16"],
        "ans": "D",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int x = 15;\nif (x % 3 == 0) {\n    if (x % 5 == 0) {\n        cout << \"A\";\n    } else {\n        cout << \"B\";\n    }\n} else {\n    cout << \"C\";\n}",
        "opts": ["A. A", "B. B", "C. C", "D. 无输出"],
        "ans": "A",
    },
    {
        "q": "下面程序片段执行后，sum 的值是？",
        "code": "int sum = 0;\nfor (int i = 1; i <= 6; i++) {\n    if (i % 3 == 0) {\n        sum += i * 2;\n    } else {\n        sum += i;\n    }\n}\ncout << sum;",
        "opts": ["A. 21", "B. 27", "C. 30", "D. 36"],
        "ans": "C",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "for (int i = 1; i <= 4; i++) {\n    for (int j = 1; j <= i; j++) {\n        cout << i;\n    }\n}",
        "opts": ["A. 1234", "B. 1121231234", "C. 1223334444", "D. 1111222233334444"],
        "ans": "C",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int cnt = 0;\nfor (int i = 1; i <= 5; i++) {\n    for (int j = i + 1; j <= 5; j++) {\n        cnt++;\n    }\n}\ncout << cnt;",
        "opts": ["A. 5", "B. 10", "C. 15", "D. 25"],
        "ans": "B",
    },
    {
        "q": "下面程序片段统计的是？",
        "code": "int cnt = 0;\nfor (int i = 1; i <= 6; i++) {\n    for (int j = 1; j <= 6; j++) {\n        if (i < j && i + j == 7) {\n            cnt++;\n        }\n    }\n}\ncout << cnt;",
        "opts": ["A. 2", "B. 3", "C. 5", "D. 6"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int s = 0;\nfor (int i = 1; i <= 4; i++) {\n    for (int j = 1; j <= 4; j++) {\n        if (i * j % 2 == 0) {\n            s++;\n        }\n    }\n}\ncout << s;",
        "opts": ["A. 8", "B. 10", "C. 12", "D. 16"],
        "ans": "C",
    },
    {
        "q": "执行一趟冒泡排序后，数组 a 的内容是？",
        "code": "int a[5] = {5, 2, 4, 1, 3};\nfor (int j = 0; j < 4; j++) {\n    if (a[j] > a[j + 1]) {\n        int t = a[j];\n        a[j] = a[j + 1];\n        a[j + 1] = t;\n    }\n}",
        "opts": ["A. {2, 4, 1, 3, 5}", "B. {2, 5, 1, 3, 4}", "C. {1, 2, 3, 4, 5}", "D. {5, 4, 3, 2, 1}"],
        "ans": "A",
    },
    {
        "q": "下面程序片段执行后，数组 a 的前 5 个元素为？",
        "code": "int a[5] = {4, 3, 1, 5, 2};\nfor (int i = 0; i < 4; i++) {\n    int k = i;\n    for (int j = i + 1; j < 5; j++) {\n        if (a[j] < a[k]) {\n            k = j;\n        }\n    }\n    int t = a[i];\n    a[i] = a[k];\n    a[k] = t;\n}",
        "opts": ["A. {4, 3, 1, 5, 2}", "B. {1, 2, 3, 4, 5}", "C. {1, 3, 4, 5, 2}", "D. {2, 1, 3, 4, 5}"],
        "ans": "B",
    },
    {
        "q": "下面插入排序片段执行到 i == 2 这一轮结束后，数组 a 为？",
        "code": "int a[5] = {3, 5, 2, 4, 1};\nfor (int i = 1; i <= 2; i++) {\n    int x = a[i];\n    int j = i - 1;\n    while (j >= 0 && a[j] > x) {\n        a[j + 1] = a[j];\n        j--;\n    }\n    a[j + 1] = x;\n}",
        "opts": ["A. {2, 3, 5, 4, 1}", "B. {3, 2, 5, 4, 1}", "C. {3, 5, 2, 4, 1}", "D. {1, 2, 3, 4, 5}"],
        "ans": "A",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int a[6] = {2, 2, 3, 3, 3, 4};\nint cnt = 0;\nfor (int i = 1; i < 6; i++) {\n    if (a[i] == a[i - 1]) {\n        cnt++;\n    }\n}\ncout << cnt;",
        "opts": ["A. 2", "B. 3", "C. 4", "D. 5"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "string s = \"314159\";\nint sum = 0;\nfor (int i = 0; i < s.size(); i++) {\n    if ((s[i] - '0') % 2 == 1) {\n        sum += s[i] - '0';\n    }\n}\ncout << sum;",
        "opts": ["A. 13", "B. 17", "C. 23", "D. 314159"],
        "ans": "C",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "string s = \"a1b2c3\";\nint x = 0;\nfor (int i = 0; i < s.size(); i++) {\n    if (s[i] >= '0' && s[i] <= '9') {\n        x = x * 10 + (s[i] - '0');\n    }\n}\ncout << x;",
        "opts": ["A. 6", "B. 123", "C. 10203", "D. 112233"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "string s = \"bcad\";\nint cnt = 0;\nfor (int i = 0; i + 1 < s.size(); i++) {\n    if (s[i] > s[i + 1]) {\n        cnt++;\n    }\n}\ncout << cnt;",
        "opts": ["A. 1", "B. 2", "C. 3", "D. 4"],
        "ans": "A",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "string s = \"120304\";\nint x = 0;\nfor (int i = 0; i < s.size(); i++) {\n    if (s[i] != '0') {\n        x += s[i] - '0';\n    }\n}\ncout << x;",
        "opts": ["A. 10", "B. 120304", "C. 1234", "D. 6"],
        "ans": "A",
    },
    {
        "q": "下面函数调用 f(6) 的返回值是？",
        "code": "int f(int n) {\n    if (n <= 1) {\n        return 1;\n    }\n    return f(n - 2) + n;\n}",
        "opts": ["A. 10", "B. 12", "C. 13", "D. 16"],
        "ans": "C",
    },
    {
        "q": "下面函数调用 g(13) 的返回值是？",
        "code": "int g(int n) {\n    if (n == 0) {\n        return 0;\n    }\n    return n % 10 + g(n / 10);\n}",
        "opts": ["A. 3", "B. 4", "C. 13", "D. 31"],
        "ans": "B",
    },
    {
        "q": "下面函数调用 h(4) 的返回值是？",
        "code": "int h(int n) {\n    if (n == 1) {\n        return 1;\n    }\n    return h(n - 1) * 2;\n}",
        "opts": ["A. 4", "B. 6", "C. 8", "D. 16"],
        "ans": "C",
    },
    {
        "q": "下面递推程序的输出结果是？",
        "code": "int a[8];\na[1] = 1;\na[2] = 3;\nfor (int i = 3; i <= 7; i++) {\n    a[i] = a[i - 1] + a[i - 2];\n}\ncout << a[7];",
        "opts": ["A. 13", "B. 18", "C. 21", "D. 29"],
        "ans": "D",
    },
    {
        "q": "下面程序片段用于判断 n 是否为质数。若 n = 49，循环中第一次使 ok 变为 false 的 i 是？",
        "code": "int n = 49;\nbool ok = true;\nfor (int i = 2; i * i <= n; i++) {\n    if (n % i == 0) {\n        ok = false;\n        break;\n    }\n}",
        "opts": ["A. 2", "B. 5", "C. 7", "D. 49"],
        "ans": "C",
    },
    {
        "q": "用筛法标记 2 到 40 的合数，最后未被标记的数有多少个？",
        "code": None,
        "opts": ["A. 10", "B. 11", "C. 12", "D. 13"],
        "ans": "C",
    },
    {
        "q": "下面筛法程序输出结果是？",
        "code": "bool vis[21] = {};\nfor (int i = 2; i <= 20; i++) {\n    if (!vis[i]) {\n        for (int j = i * i; j <= 20; j += i) {\n            vis[j] = true;\n        }\n    }\n}\nint cnt = 0;\nfor (int i = 2; i <= 20; i++) {\n    if (!vis[i]) {\n        cnt++;\n    }\n}\ncout << cnt;",
        "opts": ["A. 6", "B. 7", "C. 8", "D. 9"],
        "ans": "C",
    },
    {
        "q": "下面程序片段统计 1 到 30 中既不是 2 的倍数也不是 3 的倍数的数有多少个。输出结果是？",
        "code": "int cnt = 0;\nfor (int i = 1; i <= 30; i++) {\n    if (i % 2 != 0 && i % 3 != 0) {\n        cnt++;\n    }\n}\ncout << cnt;",
        "opts": ["A. 8", "B. 10", "C. 12", "D. 15"],
        "ans": "B",
    },
    {
        "q": "1 到 50 中能被 4 或 6 整除的整数有多少个？",
        "code": None,
        "opts": ["A. 16", "B. 17", "C. 18", "D. 20"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int ans = 0;\nfor (int i = 1; i <= 24; i++) {\n    if (i % 2 == 0 || i % 3 == 0) {\n        ans++;\n    }\n}\ncout << ans;",
        "opts": ["A. 12", "B. 16", "C. 18", "D. 20"],
        "ans": "B",
    },
    {
        "q": "用面值 1、5、10、20 的纸币凑出 48 元且张数尽量少，最少需要几张？",
        "code": None,
        "opts": ["A. 5", "B. 6", "C. 7", "D. 8"],
        "ans": "C",
    },
    {
        "q": "有任务耗时 2、2、3、5、7、9，总时间不超过 10，若要完成任务数量尽量多，最多完成几个？",
        "code": None,
        "opts": ["A. 3", "B. 4", "C. 5", "D. 6"],
        "ans": "B",
    },
    {
        "q": "有物品重量 1、4、4、5、6，背包容量为 10，只要求装入物品数量尽量多，最多能装几件？",
        "code": None,
        "opts": ["A. 2", "B. 3", "C. 4", "D. 5"],
        "ans": "B",
    },
    {
        "q": "下面二分查找程序输出结果是？",
        "code": "int a[7] = {1, 3, 4, 6, 8, 9, 11};\nint l = 0, r = 6, x = 5, ans = -1;\nwhile (l <= r) {\n    int mid = (l + r) / 2;\n    if (a[mid] == x) {\n        ans = mid;\n        break;\n    } else if (a[mid] < x) {\n        l = mid + 1;\n    } else {\n        r = mid - 1;\n    }\n}\ncout << ans;",
        "opts": ["A. -1", "B. 2", "C. 3", "D. 4"],
        "ans": "A",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int a[5] = {1, 2, 1, 2, 1};\nint ans = 0;\nfor (int i = 0; i < 5; i++) {\n    for (int j = i + 1; j < 5; j++) {\n        if (a[i] == a[j]) {\n            ans++;\n        }\n    }\n}\ncout << ans;",
        "opts": ["A. 3", "B. 4", "C. 5", "D. 6"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int a[3][3] = {\n    {1, 2, 3},\n    {4, 5, 6},\n    {7, 8, 9}\n};\nint s = 0;\nfor (int i = 0; i < 3; i++) {\n    for (int j = 0; j < 3; j++) {\n        if (i + j == 2) {\n            s += a[i][j];\n        }\n    }\n}\ncout << s;",
        "opts": ["A. 12", "B. 15", "C. 18", "D. 20"],
        "ans": "B",
    },
    {
        "q": "下面程序片段输出的是二维数组哪一部分元素之和？",
        "code": "int s = 0;\nfor (int i = 0; i < 4; i++) {\n    for (int j = 0; j < 4; j++) {\n        if (i < j) {\n            s += a[i][j];\n        }\n    }\n}",
        "opts": ["A. 主对角线", "B. 副对角线", "C. 主对角线上方", "D. 主对角线下方"],
        "ans": "C",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "bool ok = false;\nint x = 9;\nif (x % 2 == 0 || x % 3 == 0 && x % 5 == 0) {\n    ok = true;\n}\ncout << ok;",
        "opts": ["A. 0", "B. 1", "C. true", "D. false"],
        "ans": "A",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int x = 12;\nbool ok = !(x % 2 == 0 && x % 5 == 0);\ncout << ok;",
        "opts": ["A. 0", "B. 1", "C. true", "D. false"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int x = 7;\ncout << (x > 5 && x < 10 || x == 3);",
        "opts": ["A. 0", "B. 1", "C. 7", "D. 编译错误"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int a = 5, b = 2;\ncout << (a / b == 2 && a % b == 1);",
        "opts": ["A. 0", "B. 1", "C. 2", "D. 5"],
        "ans": "B",
    },
    {
        "q": "甲、乙、丙三人中只有一人说真话。甲说：“乙不是第一。”乙说：“丙是第一。”丙说：“我不是第一。”如果只有一人第一，谁是第一？",
        "code": None,
        "opts": ["A. 甲", "B. 乙", "C. 丙", "D. 无法确定"],
        "ans": "B",
    },
    {
        "q": "甲、乙、丙三人名次互不相同。甲说：“我不是第二。”乙说：“甲是第一。”丙说：“乙不是第一。”三句话中恰有两句为真，谁是第二？",
        "code": None,
        "opts": ["A. 甲", "B. 乙", "C. 丙", "D. 无法确定"],
        "ans": "A",
    },
    {
        "q": "有 4 个红球、3 个蓝球、3 个白球，随机取 1 个球，取到红球或蓝球的概率是？",
        "code": None,
        "opts": ["A. 3/10", "B. 4/10", "C. 7/10", "D. 1/2"],
        "ans": "C",
    },
    {
        "q": "从数字 1、2、3、4、5 中选出 2 个不同数字组成两位数，其中偶数有多少个？",
        "code": None,
        "opts": ["A. 6", "B. 8", "C. 10", "D. 20"],
        "ans": "B",
    },
    {
        "q": "下面枚举程序的输出结果是？",
        "code": "int cnt = 0;\nfor (int a = 1; a <= 4; a++) {\n    for (int b = 1; b <= 4; b++) {\n        for (int c = 1; c <= 4; c++) {\n            if (a != b && b != c && a != c && a + b + c == 7) {\n                cnt++;\n            }\n        }\n    }\n}\ncout << cnt;",
        "opts": ["A. 3", "B. 4", "C. 6", "D. 12"],
        "ans": "C",
    },
    {
        "q": "下面递推程序的输出结果是？",
        "code": "int a[6];\na[0] = 1;\nfor (int i = 1; i <= 5; i++) {\n    a[i] = a[i - 1] + i;\n}\ncout << a[5];",
        "opts": ["A. 11", "B. 15", "C. 16", "D. 21"],
        "ans": "C",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int n = 2026, cnt = 0;\nwhile (n > 0) {\n    if (n % 10 % 2 == 0) {\n        cnt++;\n    }\n    n /= 10;\n}\ncout << cnt;",
        "opts": ["A. 2", "B. 3", "C. 4", "D. 2026"],
        "ans": "C",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int n = 1234, x = 0;\nwhile (n > 0) {\n    int d = n % 10;\n    if (d % 2 == 0) {\n        x = x * 10 + d;\n    }\n    n /= 10;\n}\ncout << x;",
        "opts": ["A. 24", "B. 42", "C. 1234", "D. 4321"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int a = 1, b = 1;\nfor (int i = 1; i <= 4; i++) {\n    int c = a + b;\n    a = b;\n    b = c;\n}\ncout << b;",
        "opts": ["A. 5", "B. 8", "C. 13", "D. 21"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int cnt = 0;\nfor (int i = 1; i <= 5; i++) {\n    for (int j = 1; j <= 5; j++) {\n        if (i <= j && (i + j) % 2 == 0) {\n            cnt++;\n        }\n    }\n}\ncout << cnt;",
        "opts": ["A. 6", "B. 8", "C. 9", "D. 15"],
        "ans": "C",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "string s = \"cbaabc\";\nint cnt = 0;\nfor (int i = 0; i < s.size(); i++) {\n    if (s[i] == s[s.size() - 1 - i]) {\n        cnt++;\n    }\n}\ncout << cnt;",
        "opts": ["A. 0", "B. 2", "C. 4", "D. 6"],
        "ans": "D",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int a[5] = {6, 1, 6, 2, 6};\nint k = 0;\nfor (int i = 1; i < 5; i++) {\n    if (a[i] >= a[k]) {\n        k = i;\n    }\n}\ncout << k;",
        "opts": ["A. 0", "B. 2", "C. 4", "D. 6"],
        "ans": "C",
    },
    {
        "q": "下面函数调用 f(5) 的返回值是？",
        "code": "int f(int n) {\n    if (n == 0) {\n        return 0;\n    }\n    if (n % 2 == 0) {\n        return f(n - 1);\n    }\n    return f(n - 1) + n;\n}",
        "opts": ["A. 6", "B. 8", "C. 9", "D. 15"],
        "ans": "C",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int n = 18;\nint cnt = 0;\nfor (int i = 1; i * i <= n; i++) {\n    if (n % i == 0) {\n        cnt++;\n        if (i * i != n) {\n            cnt++;\n        }\n    }\n}\ncout << cnt;",
        "opts": ["A. 4", "B. 6", "C. 8", "D. 9"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int ans = 0;\nfor (int i = 10; i <= 30; i++) {\n    if (i % 4 == 0 || i % 7 == 0) {\n        ans++;\n    }\n}\ncout << ans;",
        "opts": ["A. 6", "B. 7", "C. 8", "D. 9"],
        "ans": "B",
    },
]


KEYWORDS = {
    "int", "long", "double", "float", "bool", "char", "string", "void",
    "if", "else", "for", "while", "return", "break", "continue", "true", "false",
    "cout", "cin", "include", "using", "namespace", "std",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_colored_run(paragraph, text, color="24292F", bold=False):
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    return run


TOKEN_RE = re.compile(
    r'("(?:\\.|[^"\\])*")|(\'(?:\\.|[^\'\\])*\')|(//.*)|([A-Za-z_][A-Za-z0-9_]*)|(\d+)|(==|!=|<=|>=|&&|\|\||\+\+|--|[+\-*/%=<>!?:])|(.)'
)


def add_code_line(doc, line):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.right_indent = Cm(0.25)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    set_paragraph_shading(p, "F6F8FA")
    for match in TOKEN_RE.finditer(line):
        string_dq, string_sq, comment, ident, number, op, other = match.groups()
        if string_dq or string_sq:
            add_colored_run(p, string_dq or string_sq, "A31515")
        elif comment:
            add_colored_run(p, comment, "008000")
        elif ident:
            if ident in KEYWORDS:
                add_colored_run(p, ident, "005CC5", True)
            else:
                add_colored_run(p, ident, "24292F")
        elif number:
            add_colored_run(p, number, "6F42C1")
        elif op:
            add_colored_run(p, op, "D73A49", False)
        else:
            add_colored_run(p, other, "24292F")


def add_question(doc, index, item):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{index}. {item['q']}")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.5)
    r.bold = True

    if item.get("code"):
        for line in item["code"].split("\n"):
            add_code_line(doc, line)
        doc.add_paragraph().paragraph_format.space_after = Pt(1)

    for opt in item["opts"]:
        po = doc.add_paragraph()
        po.paragraph_format.left_indent = Cm(0.55)
        po.paragraph_format.space_after = Pt(1)
        ro = po.add_run(opt)
        ro.font.name = "Microsoft YaHei"
        ro._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        ro.font.size = Pt(10)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("小学组提高级 C++ 错题薄弱点专项训练（50题）")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(16)
    run.bold = True

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("针对错题知识点：表达式、嵌套循环、排序、字符串、递归递推、筛法、贪心枚举、逻辑判断")
    nr.font.name = "Microsoft YaHei"
    nr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(89, 99, 110)

    for i, q in enumerate(questions, 1):
        add_question(doc, i, q)

    doc.add_page_break()
    ans_title = doc.add_paragraph()
    ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = ans_title.add_run("答案")
    ar.font.name = "Microsoft YaHei"
    ar._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    ar.font.size = Pt(16)
    ar.bold = True

    table = doc.add_table(rows=11, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["1-10", "11-20", "21-30", "31-40", "41-50"]
    for c, h in enumerate(headers):
        cell = table.cell(0, c)
        set_cell_shading(cell, "DDEBFF")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10)

    for row in range(1, 11):
        for col in range(5):
            num = row + col * 10
            cell = table.cell(row, col)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text = f"{num}. {questions[num - 1]['ans']}"
            r = p.add_run(text)
            r.font.name = "Microsoft YaHei"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            r.font.size = Pt(10.5)

    doc.save(OUT)
    print(OUT)
    print(f"questions={len(questions)} code_questions={sum(1 for q in questions if q.get('code'))}")


if __name__ == "__main__":
    build_doc()
