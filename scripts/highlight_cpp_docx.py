import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SRC = Path("/Users/cen/Downloads/20道题01_50题含答案.docx")
OUT = Path("/Users/cen/Documents/Codex/2026-05-18/4-5-1-2-3-4/20道题01_50题含答案_C++语法高亮.docx")


KEYWORDS = {
    "int", "long", "double", "float", "bool", "char", "string", "void",
    "if", "else", "for", "while", "return", "break", "continue",
}
CONSTANTS = {"true", "false", "nullptr", "NULL"}
IO_WORDS = {"cin", "cout", "endl"}

TOKEN_RE = re.compile(
    r'//.*?$'
    r'|"(?:\\.|[^"\\])*"'
    r"|'(?:\\.|[^'\\])*'"
    r"|\b\d+(?:\.\d+)?\b"
    r"|\b[A-Za-z_]\w*\b"
    r"|==|!=|<=|>=|\+\+|--|\+=|-=|\*=|/=|%=|&&|\|\||<<|>>"
    r"|[{}\[\]();,?:+\-*/%<>=!&|.]",
    re.MULTILINE,
)


def set_paragraph_shading(paragraph, fill="F6F8FA") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, name="Consolas", size=9.5) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def style_token(run, token: str) -> None:
    set_run_font(run)
    run.font.color.rgb = RGBColor(36, 41, 47)

    if token.startswith("//"):
        run.font.color.rgb = RGBColor(106, 115, 125)
        run.italic = True
    elif token.startswith('"') or token.startswith("'"):
        run.font.color.rgb = RGBColor(163, 21, 21)
    elif re.fullmatch(r"\d+(?:\.\d+)?", token):
        run.font.color.rgb = RGBColor(111, 66, 193)
    elif token in KEYWORDS:
        run.font.color.rgb = RGBColor(0, 92, 197)
        run.bold = True
    elif token in CONSTANTS:
        run.font.color.rgb = RGBColor(128, 0, 128)
        run.bold = True
    elif token in IO_WORDS:
        run.font.color.rgb = RGBColor(0, 128, 128)
        run.bold = True
    elif re.fullmatch(r"==|!=|<=|>=|\+\+|--|\+=|-=|\*=|/=|%=|&&|\|\||<<|>>|[+\-*/%<>=!&|?:]", token):
        run.font.color.rgb = RGBColor(89, 99, 110)


def looks_like_cpp(text: str) -> bool:
    s = text.strip()
    if not s or s == "答案":
        return False
    if re.match(r"^\d+\. ", s):
        return False
    if re.match(r"^[A-D]\. ", s):
        body = s[3:].strip()
        return bool(re.search(r"\b(for|while|if|int|long|double|float|bool|char|string|return|cout|cin)\b|[();{}]|==|<=|>=|&&|\|\||\+\+", body))
    if re.search(r"\b(int|long|double|float|bool|char|string|void|for|while|if|else|return|cout|cin)\b", s):
        return True
    if any(mark in s for mark in [";", "{", "}", "++", "--", "+=", "-=", "*=", "/=", "%=", "==", "!=", "<=", ">=", "&&", "||", "<<", ">>"]):
        return True
    if re.fullmatch(r"\{?[0-9,\s]+\}?,?", s):
        return True
    return False


def clear_and_highlight(paragraph, text: str, option_prefix: str = "") -> None:
    paragraph.clear()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    set_paragraph_shading(paragraph)

    if option_prefix:
        prefix_run = paragraph.add_run(option_prefix)
        prefix_run.font.name = "Arial"
        prefix_run.font.size = Pt(10.5)

    pos = 0
    for match in TOKEN_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run)
            run.font.color.rgb = RGBColor(36, 41, 47)
        token = match.group(0)
        run = paragraph.add_run(token)
        style_token(run, token)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run)
        run.font.color.rgb = RGBColor(36, 41, 47)


def main() -> None:
    doc = Document(SRC)
    highlighted = 0
    for paragraph in doc.paragraphs:
        original = paragraph.text
        text = original.strip()
        if not looks_like_cpp(text):
            continue
        option_match = re.match(r"^([A-D]\.\s+)(.*)$", text)
        if option_match:
            clear_and_highlight(paragraph, option_match.group(2), option_match.group(1))
        else:
            clear_and_highlight(paragraph, text)
        highlighted += 1
    doc.save(OUT)
    print(f"{OUT}\nHighlighted paragraphs: {highlighted}")


if __name__ == "__main__":
    main()
