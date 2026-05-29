from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from make_gesp_style_paper import (
    add_heading,
    add_label_para,
    add_tests_table,
    mc_questions,
    programs,
    set_font,
)
from make_remedial_50_cpp_doc import add_question, set_cell_shading


CHOICE_OUT = "/Users/cen/Documents/Codex/2026-05-18/4-5-1-2-3-4/GESP四级偏上五级入门_选择题.docx"
PROGRAM_OUT = "/Users/cen/Documents/Codex/2026-05-18/4-5-1-2-3-4/GESP四级偏上五级入门_编程题.docx"


def setup_doc(title_text, note_text):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(title_text)
    set_font(r, 16, True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(note_text)
    set_font(nr, 9, False, "59636E")
    return doc


def build_choice_doc():
    doc = setup_doc(
        "GESP C++ 四级偏上 / 五级入门 选择题",
        "依据 GESP 官方近年 C++ 四级、五级真题考点改编：10 道选择题",
    )
    add_heading(doc, "选择题（每题 3 分，共 30 分）")
    for i, q in enumerate(mc_questions, 1):
        add_question(doc, i, q)

    doc.add_page_break()
    add_heading(doc, "答案")
    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(10):
        row = i // 5
        col = i % 5
        cell = table.cell(row, col)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if row == 0:
            set_cell_shading(cell, "F6F8FA")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{i + 1}. {mc_questions[i]['ans']}")
        set_font(r, 10.5, True)

    doc.save(CHOICE_OUT)
    print(CHOICE_OUT)


def build_program_doc():
    doc = setup_doc(
        "GESP C++ 四级偏上 / 五级入门 编程题",
        "依据 GESP 官方近年 C++ 四级、五级真题考点改编：4 道编程题，每题 10 组测试例",
    )
    add_heading(doc, "编程题（每题 25 分，共 100 分）")
    for idx, item in enumerate(programs, 1):
        add_heading(doc, f"编程题 {idx}. {item['title']}", 12)
        add_label_para(doc, "题目描述：", item["background"])
        add_label_para(doc, "输入格式：", item["input"])
        add_label_para(doc, "输出格式：", item["output"])
        add_label_para(doc, "数据范围：", item["constraints"])
        add_label_para(doc, "测试例：", "")
        add_tests_table(doc, item["tests"])

    doc.add_page_break()
    add_heading(doc, "说明")
    p = doc.add_paragraph()
    r = p.add_run("编程题参考答案以题面测试例输出为准；正式评分建议加入隐藏测试。")
    set_font(r, 10.5)

    doc.save(PROGRAM_OUT)
    print(PROGRAM_OUT)


if __name__ == "__main__":
    build_choice_doc()
    build_program_doc()
