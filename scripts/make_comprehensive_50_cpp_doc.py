from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from make_remedial_50_cpp_doc import add_question, set_cell_shading


OUT = "/Users/cen/Documents/Codex/2026-05-18/4-5-1-2-3-4/小学组提高级C++综合训练50题_高亮缩进含答案.docx"


questions = [
    {"q": "下面哪个变量名在 C++ 中是合法的？", "code": None, "opts": ["A. 3sum", "B. total-score", "C. _answer", "D. long"], "ans": "C"},
    {"q": "下面程序片段执行后，输出结果是？", "code": "int a = 4;\nint b = 7;\na = b - a;\nb = a + b;\ncout << a << \" \" << b;", "opts": ["A. 3 10", "B. 4 7", "C. 7 3", "D. 3 7"], "ans": "A"},
    {"q": "下面哪个类型最适合保存一个学生是否通过考试？", "code": None, "opts": ["A. int", "B. double", "C. bool", "D. char"], "ans": "C"},
    {"q": "执行下面程序片段后，输出结果是？", "code": "int x = 17;\ncout << x / 5 << \" \" << x % 5;", "opts": ["A. 3 2", "B. 2 3", "C. 3.4 0", "D. 17 5"], "ans": "A"},
    {"q": "若输入为 6 和 8，下面程序输出结果是？", "code": "int a, b;\ncin >> a >> b;\ncout << a * b - a;", "opts": ["A. 14", "B. 42", "C. 48", "D. 56"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "double x = 7.0 / 2;\ncout << x;", "opts": ["A. 3", "B. 3.5", "C. 4", "D. 7"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "cout << max(12, min(9, 15));", "opts": ["A. 9", "B. 12", "C. 15", "D. 21"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "cout << abs(-8) + sqrt(16);", "opts": ["A. 8", "B. 10", "C. 12", "D. 24"], "ans": "C"},
    {"q": "下面表达式的值是？", "code": "3 + 4 * 2 - 10 / 5", "opts": ["A. 7", "B. 9", "C. 11", "D. 13"], "ans": "B"},
    {"q": "下面程序片段执行后，输出结果是？", "code": "int x = 5;\nx += 3;\nx *= 2;\ncout << x;", "opts": ["A. 8", "B. 10", "C. 13", "D. 16"], "ans": "D"},
    {"q": "下面程序片段的输出结果是？", "code": "int score = 86;\nif (score >= 90) {\n    cout << \"A\";\n} else if (score >= 80) {\n    cout << \"B\";\n} else {\n    cout << \"C\";\n}", "opts": ["A. A", "B. B", "C. C", "D. 无输出"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 13;\nif (x % 2 == 0) {\n    cout << \"even\";\n} else {\n    cout << \"odd\";\n}", "opts": ["A. even", "B. odd", "C. 13", "D. 编译错误"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int a = 5, b = 9;\nint c = (a > b ? a : b);\ncout << c;", "opts": ["A. 5", "B. 9", "C. 14", "D. 0"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 12;\ncout << (x % 3 == 0 && x % 4 == 0);", "opts": ["A. 0", "B. 1", "C. 3", "D. 4"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int sum = 0;\nfor (int i = 1; i <= 5; i++) {\n    sum += i;\n}\ncout << sum;", "opts": ["A. 5", "B. 10", "C. 15", "D. 20"], "ans": "C"},
    {"q": "下面程序片段的输出结果是？", "code": "int i = 1;\nwhile (i <= 4) {\n    cout << i;\n    i += 2;\n}", "opts": ["A. 1234", "B. 13", "C. 24", "D. 135"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int cnt = 0;\nfor (int i = 1; i <= 4; i++) {\n    for (int j = 1; j <= i; j++) {\n        cnt++;\n    }\n}\ncout << cnt;", "opts": ["A. 4", "B. 8", "C. 10", "D. 16"], "ans": "C"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[5] = {2, 4, 6, 8, 10};\ncout << a[1] + a[3];", "opts": ["A. 6", "B. 10", "C. 12", "D. 14"], "ans": "C"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[6] = {3, 1, 4, 1, 5, 9};\nint mx = a[0];\nfor (int i = 1; i < 6; i++) {\n    if (a[i] > mx) {\n        mx = a[i];\n    }\n}\ncout << mx;", "opts": ["A. 1", "B. 4", "C. 5", "D. 9"], "ans": "D"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[5] = {5, 2, 5, 3, 5};\nint cnt = 0;\nfor (int i = 0; i < 5; i++) {\n    if (a[i] == 5) {\n        cnt++;\n    }\n}\ncout << cnt;", "opts": ["A. 2", "B. 3", "C. 4", "D. 5"], "ans": "B"},
    {"q": "下面二维数组程序的输出结果是？", "code": "int a[2][3] = {\n    {1, 2, 3},\n    {4, 5, 6}\n};\ncout << a[0][2] + a[1][1];", "opts": ["A. 5", "B. 7", "C. 8", "D. 9"], "ans": "C"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[3][3] = {\n    {1, 2, 3},\n    {4, 5, 6},\n    {7, 8, 9}\n};\nint s = 0;\nfor (int i = 0; i < 3; i++) {\n    s += a[i][i];\n}\ncout << s;", "opts": ["A. 12", "B. 15", "C. 18", "D. 24"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "string s = \"contest\";\ncout << s[0] << s[3];", "opts": ["A. co", "B. ct", "C. ce", "D. ns"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "string s = \"a1b2c3\";\nint sum = 0;\nfor (int i = 0; i < s.size(); i++) {\n    if (s[i] >= '0' && s[i] <= '9') {\n        sum += s[i] - '0';\n    }\n}\ncout << sum;", "opts": ["A. 3", "B. 6", "C. 123", "D. 112233"], "ans": "B"},
    {"q": "下面函数调用 add(4, 7) 的返回值是？", "code": "int add(int x, int y) {\n    return x + y;\n}", "opts": ["A. 3", "B. 4", "C. 7", "D. 11"], "ans": "D"},
    {"q": "下面函数调用 f(5) 的返回值是？", "code": "int f(int n) {\n    if (n == 1) {\n        return 1;\n    }\n    return n + f(n - 1);\n}", "opts": ["A. 5", "B. 10", "C. 15", "D. 120"], "ans": "C"},
    {"q": "下面递推程序的输出结果是？", "code": "int a[6];\na[1] = 1;\na[2] = 2;\nfor (int i = 3; i <= 5; i++) {\n    a[i] = a[i - 1] + a[i - 2];\n}\ncout << a[5];", "opts": ["A. 3", "B. 5", "C. 8", "D. 13"], "ans": "C"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[5] = {4, 1, 3, 2, 5};\nfor (int j = 0; j < 4; j++) {\n    if (a[j] > a[j + 1]) {\n        int t = a[j];\n        a[j] = a[j + 1];\n        a[j + 1] = t;\n    }\n}\ncout << a[0] << a[4];", "opts": ["A. 15", "B. 41", "C. 45", "D. 14"], "ans": "A"},
    {"q": "下面选择排序执行完 i == 0 这一轮后，数组 a 为？", "code": "int a[5] = {6, 4, 2, 5, 3};\nint k = 0;\nfor (int j = 1; j < 5; j++) {\n    if (a[j] < a[k]) {\n        k = j;\n    }\n}\nint t = a[0];\na[0] = a[k];\na[k] = t;", "opts": ["A. {2, 4, 6, 5, 3}", "B. {3, 4, 2, 5, 6}", "C. {6, 4, 2, 5, 3}", "D. {2, 3, 4, 5, 6}"], "ans": "A"},
    {"q": "下面二分查找程序的输出结果是？", "code": "int a[6] = {1, 3, 5, 7, 9, 11};\nint l = 0, r = 5, x = 7, ans = -1;\nwhile (l <= r) {\n    int mid = (l + r) / 2;\n    if (a[mid] == x) {\n        ans = mid;\n        break;\n    } else if (a[mid] < x) {\n        l = mid + 1;\n    } else {\n        r = mid - 1;\n    }\n}\ncout << ans;", "opts": ["A. 2", "B. 3", "C. 4", "D. -1"], "ans": "B"},
    {"q": "下面枚举程序的输出结果是？", "code": "int cnt = 0;\nfor (int a = 1; a <= 3; a++) {\n    for (int b = 1; b <= 3; b++) {\n        if (a != b) {\n            cnt++;\n        }\n    }\n}\ncout << cnt;", "opts": ["A. 3", "B. 6", "C. 9", "D. 12"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int cnt = 0;\nfor (int i = 1; i <= 20; i++) {\n    if (i % 2 == 0 || i % 5 == 0) {\n        cnt++;\n    }\n}\ncout << cnt;", "opts": ["A. 10", "B. 12", "C. 14", "D. 16"], "ans": "B"},
    {"q": "用筛法标记 2 到 30 的合数，最后未被标记的数有多少个？", "code": None, "opts": ["A. 8", "B. 9", "C. 10", "D. 11"], "ans": "C"},
    {"q": "下面程序片段判断 n 是否为质数。若 n = 37，最终 ok 的值是？", "code": "int n = 37;\nbool ok = true;\nfor (int i = 2; i * i <= n; i++) {\n    if (n % i == 0) {\n        ok = false;\n        break;\n    }\n}\ncout << ok;", "opts": ["A. 0", "B. 1", "C. 37", "D. 编译错误"], "ans": "B"},
    {"q": "用面值 1、5、10、20 的纸币凑出 38 元且张数尽量少，最少需要几张？", "code": None, "opts": ["A. 4", "B. 5", "C. 6", "D. 7"], "ans": "C"},
    {"q": "有任务耗时 1、2、4、6、9，总时间不超过 8，为了完成尽量多的任务，最多完成几个？", "code": None, "opts": ["A. 2", "B. 3", "C. 4", "D. 5"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[4] = {1, 2, 3, 4};\nint ans = 0;\nfor (int i = 0; i < 4; i++) {\n    for (int j = i; j < 4; j++) {\n        ans += a[i];\n    }\n}\ncout << ans;", "opts": ["A. 10", "B. 20", "C. 30", "D. 40"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int n = 1234;\nint x = 0;\nwhile (n > 0) {\n    x = x * 10 + n % 10;\n    n /= 10;\n}\ncout << x;", "opts": ["A. 1234", "B. 4321", "C. 10", "D. 1"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int a = 0, b = 1;\nfor (int i = 1; i <= 4; i++) {\n    int c = a + b;\n    a = b;\n    b = c;\n}\ncout << b;", "opts": ["A. 3", "B. 5", "C. 8", "D. 13"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[5] = {1, 2, 1, 2, 1};\nint ans = 0;\nfor (int i = 0; i < 5; i++) {\n    for (int j = i + 1; j < 5; j++) {\n        if (a[i] == a[j]) {\n            ans++;\n        }\n    }\n}\ncout << ans;", "opts": ["A. 3", "B. 4", "C. 5", "D. 6"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 9;\ncout << (x % 3 == 0 && x % 2 == 0 || x > 8);", "opts": ["A. 0", "B. 1", "C. 9", "D. 编译错误"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 10;\ncout << (x > 5 || x / 0 == 1);", "opts": ["A. 0", "B. 1", "C. 运行时错误", "D. 编译错误"], "ans": "B"},
    {"q": "甲、乙、丙三人中只有一人说真话。甲说：“乙第一。”乙说：“丙第一。”丙说：“我不是第一。”如果只有一人第一，谁第一？", "code": None, "opts": ["A. 甲", "B. 乙", "C. 丙", "D. 无法确定"], "ans": "B"},
    {"q": "有 3 个红球、2 个蓝球、5 个白球，随机取 1 个球，取到红球或蓝球的概率是？", "code": None, "opts": ["A. 1/5", "B. 1/2", "C. 3/5", "D. 4/5"], "ans": "B"},
    {"q": "从数字 1、2、3、4 中选出两个不同数字组成两位数，其中偶数有多少个？", "code": None, "opts": ["A. 4", "B. 6", "C. 8", "D. 12"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int cnt = 0;\nfor (int a = 1; a <= 4; a++) {\n    for (int b = 1; b <= 4; b++) {\n        if (a < b && a + b == 5) {\n            cnt++;\n        }\n    }\n}\ncout << cnt;", "opts": ["A. 1", "B. 2", "C. 3", "D. 4"], "ans": "B"},
    {"q": "下面程序片段模拟的是哪种操作？", "code": "int x = 0;\nfor (int i = 1; i <= n; i++) {\n    x += i;\n}", "opts": ["A. 求 1 到 n 的和", "B. 求 n 的阶乘", "C. 判断 n 是否为质数", "D. 数组排序"], "ans": "A"},
    {"q": "下面程序片段的输出结果是？", "code": "int a = 2, b = 3, c = 4;\ncout << (a + b * c);", "opts": ["A. 14", "B. 20", "C. 24", "D. 36"], "ans": "A"},
    {"q": "下面程序片段的输出结果是？", "code": "char ch = 'A';\ncout << ch;", "opts": ["A. A", "B. ch", "C. 65", "D. 编译错误"], "ans": "A"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 4;\ncout << ((x > 3 ? x + 1 : x - 1) == 5);", "opts": ["A. 0", "B. 1", "C. 4", "D. 5"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 24;\nint cnt = 0;\nfor (int i = 1; i * i <= x; i++) {\n    if (x % i == 0) {\n        cnt++;\n        if (i * i != x) {\n            cnt++;\n        }\n    }\n}\ncout << cnt;", "opts": ["A. 6", "B. 8", "C. 10", "D. 12"], "ans": "B"},
]


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("小学组提高级 C++ 综合训练（50题）")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(16)
    run.bold = True

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("覆盖变量、输入输出、运算、分支循环、数组字符串、函数递归、算法、概率统计与逻辑推理")
    nr.font.name = "Microsoft YaHei"
    nr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(89, 99, 110)

    for i, q in enumerate(questions, 1):
        add_question(doc, i, q)

    doc.add_page_break()
    ans_title = doc.add_paragraph()
    ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = ans_title.add_run("答案")
    ar.font.name = "Microsoft YaHei"
    ar._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    ar.font.size = Pt(16)
    ar.bold = True

    table = doc.add_table(rows=11, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["1-10", "11-20", "21-30", "31-40", "41-50"]
    for c, h in enumerate(headers):
        cell = table.cell(0, c)
        set_cell_shading(cell, "DDEBFF")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10)

    for row in range(1, 11):
        for col in range(5):
            num = row + col * 10
            cell = table.cell(row, col)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"{num}. {questions[num - 1]['ans']}")
            r.font.name = "Microsoft YaHei"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            r.font.size = Pt(10.5)

    doc.save(OUT)
    print(OUT)
    print(f"questions={len(questions)} code_questions={sum(1 for q in questions if q.get('code'))}")


if __name__ == "__main__":
    build_doc()
