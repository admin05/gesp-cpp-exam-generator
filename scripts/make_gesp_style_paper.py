from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from make_remedial_50_cpp_doc import add_question, set_cell_shading, add_code_line


OUT = "/Users/cen/Documents/Codex/2026-05-18/4-5-1-2-3-4/GESP四级偏上五级入门真题风格卷_10选4编程.docx"


mc_questions = [
    {
        "q": "下面程序片段的输出结果是？",
        "code": "char ch = '7';\nint x = ch - '0';\ncout << x * 3;",
        "opts": ["A. 21", "B. 777", "C. 165", "D. 编译错误"],
        "ans": "A",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int a = 5, b = 2, c = 3;\ncout << a + b * c - a / b;",
        "opts": ["A. 9", "B. 10", "C. 11", "D. 13"],
        "ans": "A",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int cnt = 0;\nfor (int i = 1; i <= 4; i++) {\n    for (int j = i; j <= 4; j++) {\n        if ((i + j) % 2 == 0) {\n            cnt++;\n        }\n    }\n}\ncout << cnt;",
        "opts": ["A. 4", "B. 6", "C. 8", "D. 10"],
        "ans": "B",
    },
    {
        "q": "执行下面一趟冒泡排序后，数组 a 的内容是？",
        "code": "int a[5] = {4, 1, 5, 2, 3};\nfor (int j = 0; j < 4; j++) {\n    if (a[j] > a[j + 1]) {\n        int t = a[j];\n        a[j] = a[j + 1];\n        a[j + 1] = t;\n    }\n}",
        "opts": ["A. {1, 4, 2, 3, 5}", "B. {1, 2, 3, 4, 5}", "C. {4, 1, 2, 3, 5}", "D. {1, 4, 5, 2, 3}"],
        "ans": "A",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int a[3][3] = {\n    {2, 4, 6},\n    {1, 3, 5},\n    {7, 9, 8}\n};\nint s = 0;\nfor (int i = 0; i < 3; i++) {\n    s += a[i][2 - i];\n}\ncout << s;",
        "opts": ["A. 12", "B. 16", "C. 18", "D. 20"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "string s = \"a12b03\";\nint x = 0;\nfor (int i = 0; i < s.size(); i++) {\n    if (s[i] >= '0' && s[i] <= '9') {\n        x = x * 10 + (s[i] - '0');\n    }\n}\ncout << x;",
        "opts": ["A. 15", "B. 1203", "C. 123", "D. 10203"],
        "ans": "B",
    },
    {
        "q": "下面函数调用 f(5) 的返回值是？",
        "code": "int f(int n) {\n    if (n <= 1) {\n        return 1;\n    }\n    return f(n - 1) + 2 * f(n - 2);\n}",
        "opts": ["A. 11", "B. 15", "C. 21", "D. 31"],
        "ans": "C",
    },
    {
        "q": "下面二分查找程序中，第一次循环结束后 l 和 r 的值分别是？",
        "code": "int a[7] = {1, 3, 5, 7, 9, 11, 13};\nint l = 0, r = 6, x = 6;\nint mid = (l + r) / 2;\nif (a[mid] < x) {\n    l = mid + 1;\n} else {\n    r = mid - 1;\n}",
        "opts": ["A. l = 0, r = 2", "B. l = 4, r = 6", "C. l = 0, r = 6", "D. l = 3, r = 6"],
        "ans": "A",
    },
    {
        "q": "用筛法标记 2 到 50 的合数，最后未被标记的数有多少个？",
        "code": None,
        "opts": ["A. 13", "B. 14", "C. 15", "D. 16"],
        "ans": "C",
    },
    {
        "q": "下面逻辑表达式的值是？",
        "code": "int x = 20;\ncout << (x % 4 == 0 && x % 6 == 0 || x % 10 == 0);",
        "opts": ["A. 0", "B. 1", "C. 4", "D. 10"],
        "ans": "B",
    },
]


programs = [
    {
        "title": "矩阵十字和",
        "background": "给定一个 n 行 m 列的整数矩阵，以及一个位置 (x, y)。请输出第 x 行与第 y 列所有元素之和，其中交叉位置只计算一次。行列编号从 1 开始。",
        "input": "第一行输入 n 和 m。接下来 n 行，每行 m 个整数。最后一行输入 x 和 y。",
        "output": "输出一个整数，表示矩阵十字和。",
        "constraints": "1 <= n, m <= 50；矩阵元素为 0 到 1000 的整数。",
        "tests": [
            ("3 3\n1 2 3\n4 5 6\n7 8 9\n2 2", "25"),
            ("2 4\n1 1 1 1\n2 2 2 2\n1 3", "6"),
            ("4 2\n3 5\n7 9\n1 4\n6 8\n4 1", "28"),
            ("1 5\n2 4 6 8 10\n1 4", "30"),
            ("5 1\n1\n2\n3\n4\n5\n3 1", "15"),
            ("2 2\n10 20\n30 40\n2 2", "90"),
            ("3 4\n1 3 5 7\n2 4 6 8\n9 11 13 15\n3 1", "58"),
            ("4 4\n1 2 3 4\n5 6 7 8\n9 10 11 12\n13 14 15 16\n1 1", "28"),
            ("3 2\n100 1\n2 3\n4 5\n1 2", "109"),
            ("2 3\n7 8 9\n1 2 3\n2 1", "13"),
        ],
    },
    {
        "title": "区间质数个数",
        "background": "给定两个整数 L 和 R，请统计闭区间 [L, R] 中质数的个数。",
        "input": "输入一行，包含两个整数 L 和 R。",
        "output": "输出一个整数，表示区间内质数的个数。",
        "constraints": "1 <= L <= R <= 100000。",
        "tests": [
            ("1 10", "4"),
            ("2 2", "1"),
            ("14 20", "2"),
            ("30 50", "5"),
            ("90 100", "1"),
            ("101 120", "5"),
            ("1 1", "0"),
            ("997 1009", "3"),
            ("1000 1050", "8"),
            ("99990 100000", "0"),
        ],
    },
    {
        "title": "最多完成任务数",
        "background": "有 n 个任务，每个任务需要一定时间。小杨总时间为 T，每个任务最多完成一次。为了完成尽量多的任务，请输出最多能完成几个任务。",
        "input": "第一行输入 n 和 T。第二行输入 n 个正整数，表示每个任务需要的时间。",
        "output": "输出一个整数，表示最多能完成的任务数量。",
        "constraints": "1 <= n <= 1000；1 <= T <= 100000；每个任务耗时为 1 到 100000。",
        "tests": [
            ("5 10\n1 3 4 6 8", "3"),
            ("4 5\n2 2 2 2", "2"),
            ("6 12\n5 1 2 7 3 4", "4"),
            ("3 100\n30 40 50", "2"),
            ("5 7\n8 9 10 11 12", "0"),
            ("1 5\n5", "1"),
            ("7 15\n1 1 2 2 3 8 9", "5"),
            ("6 20\n10 5 5 4 6 2", "4"),
            ("8 9\n4 4 4 1 1 1 10 2", "5"),
            ("5 6\n6 1 2 3 4", "3"),
        ],
    },
    {
        "title": "递推数列",
        "background": "定义数列 f(1)=1，f(2)=2。当 n>=3 时，f(n)=f(n-1)+2*f(n-2)。给定 n，请输出 f(n)。",
        "input": "输入一个整数 n。",
        "output": "输出 f(n) 的值。",
        "constraints": "1 <= n <= 40，答案在 long long 范围内。",
        "tests": [
            ("1", "1"),
            ("2", "2"),
            ("3", "4"),
            ("4", "8"),
            ("5", "16"),
            ("6", "32"),
            ("7", "64"),
            ("10", "512"),
            ("20", "524288"),
            ("30", "536870912"),
        ],
    },
]


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level_size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, level_size, True)
    return p


def add_label_para(doc, label, body):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.25)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label)
    set_font(r1, 10.5, True)
    r2 = p.add_run(body)
    set_font(r2, 10.5)


def add_io_block(doc, text):
    for line in text.split("\n"):
        add_code_line(doc, line)


def add_tests_table(doc, tests):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["序号", "输入", "输出"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "DDEBFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, 9.5, True)
    for idx, (inp, out) in enumerate(tests, 1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = inp
        cells[2].text = out
        for c in cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = "Courier New" if c in (cells[1], cells[2]) else "Microsoft YaHei"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), r.font.name)
                    r.font.size = Pt(8.5)
    doc.add_paragraph()


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("GESP C++ 四级偏上 / 五级入门 真题风格模拟卷")
    set_font(r, 16, True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("依据 GESP 官方近年 C++ 四级、五级真题考点改编：10 道选择题 + 4 道编程题")
    set_font(nr, 9, False, "59636E")

    add_heading(doc, "一、选择题（每题 3 分，共 30 分）")
    for i, q in enumerate(mc_questions, 1):
        add_question(doc, i, q)

    add_heading(doc, "二、编程题（每题 25 分，共 100 分）")
    for idx, item in enumerate(programs, 1):
        add_heading(doc, f"编程题 {idx}. {item['title']}", 12)
        add_label_para(doc, "题目描述：", item["background"])
        add_label_para(doc, "输入格式：", item["input"])
        add_label_para(doc, "输出格式：", item["output"])
        add_label_para(doc, "数据范围：", item["constraints"])
        add_label_para(doc, "测试例：", "")
        add_tests_table(doc, item["tests"])

    doc.add_page_break()
    add_heading(doc, "答案与参考来源")
    table = doc.add_table(rows=3, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(10):
        row = i // 5
        col = i % 5
        cell = table.cell(row, col)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{i + 1}. {mc_questions[i]['ans']}")
        set_font(r, 10.5, True)

    p = doc.add_paragraph()
    r = p.add_run("编程题参考答案：以题面测试例输出为准；建议评分时同时加入隐藏测试。")
    set_font(r, 10.5)

    refs = [
        "GESP 官网：真题及解析栏目（2023-2026 年多次认证，含 C++ 四级、五级试题入口）",
        "GESP 官网：2025年3月、2025年6月、2024年12月、2024年9月、2023年12月认证真题页面",
        "说明：本卷为依据官方真题考点改编的模拟卷，不逐字复制官方试题。",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        r = p.add_run(ref)
        set_font(r, 9.5)

    doc.save(OUT)
    print(OUT)
    print(f"mc={len(mc_questions)} programs={len(programs)} tests={sum(len(p['tests']) for p in programs)}")


if __name__ == "__main__":
    build_doc()
