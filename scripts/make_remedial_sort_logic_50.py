from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from make_remedial_50_cpp_doc import add_question, set_cell_shading


OUT = "/Users/cen/Documents/Codex/2026-05-18/4-5-1-2-3-4/排序数组整除逻辑专项训练50题_C++高亮缩进.docx"


questions = [
    {"q": "执行下面一趟冒泡排序后，数组 a 的内容是？", "code": "int a[5] = {6, 2, 5, 1, 4};\nfor (int j = 0; j < 4; j++) {\n    if (a[j] > a[j + 1]) {\n        int t = a[j];\n        a[j] = a[j + 1];\n        a[j + 1] = t;\n    }\n}", "opts": ["A. {2, 5, 1, 4, 6}", "B. {1, 2, 4, 5, 6}", "C. {2, 6, 1, 4, 5}", "D. {6, 5, 2, 4, 1}"], "ans": "A"},
    {"q": "下面冒泡排序执行完外层 i == 1 这一轮后，数组 a 为？", "code": "int a[5] = {4, 1, 3, 5, 2};\nfor (int i = 0; i <= 1; i++) {\n    for (int j = 0; j < 4 - i; j++) {\n        if (a[j] > a[j + 1]) {\n            int t = a[j];\n            a[j] = a[j + 1];\n            a[j + 1] = t;\n        }\n    }\n}", "opts": ["A. {1, 3, 4, 2, 5}", "B. {1, 3, 2, 4, 5}", "C. {1, 2, 3, 4, 5}", "D. {4, 1, 3, 2, 5}"], "ans": "B"},
    {"q": "下面选择排序执行完 i == 0 这一轮后，数组 a 为？", "code": "int a[5] = {7, 3, 6, 2, 5};\nint k = 0;\nfor (int j = 1; j < 5; j++) {\n    if (a[j] < a[k]) {\n        k = j;\n    }\n}\nint t = a[0];\na[0] = a[k];\na[k] = t;", "opts": ["A. {2, 3, 6, 7, 5}", "B. {3, 7, 6, 2, 5}", "C. {2, 3, 5, 6, 7}", "D. {7, 3, 6, 2, 5}"], "ans": "A"},
    {"q": "下面选择排序执行完 i == 1 这一轮后，数组 a 为？", "code": "int a[5] = {5, 4, 1, 3, 2};\nfor (int i = 0; i <= 1; i++) {\n    int k = i;\n    for (int j = i + 1; j < 5; j++) {\n        if (a[j] < a[k]) {\n            k = j;\n        }\n    }\n    int t = a[i];\n    a[i] = a[k];\n    a[k] = t;\n}", "opts": ["A. {1, 2, 5, 3, 4}", "B. {1, 4, 5, 3, 2}", "C. {1, 2, 3, 4, 5}", "D. {2, 1, 5, 3, 4}"], "ans": "A"},
    {"q": "下面插入排序执行到 i == 2 这一轮结束后，数组 a 为？", "code": "int a[5] = {6, 2, 4, 1, 5};\nfor (int i = 1; i <= 2; i++) {\n    int x = a[i];\n    int j = i - 1;\n    while (j >= 0 && a[j] > x) {\n        a[j + 1] = a[j];\n        j--;\n    }\n    a[j + 1] = x;\n}", "opts": ["A. {2, 4, 6, 1, 5}", "B. {2, 6, 4, 1, 5}", "C. {1, 2, 4, 6, 5}", "D. {6, 2, 4, 1, 5}"], "ans": "A"},
    {"q": "下面插入排序执行完全部循环后，数组 a 为？", "code": "int a[5] = {3, 1, 4, 2, 5};\nfor (int i = 1; i < 5; i++) {\n    int x = a[i];\n    int j = i - 1;\n    while (j >= 0 && a[j] > x) {\n        a[j + 1] = a[j];\n        j--;\n    }\n    a[j + 1] = x;\n}", "opts": ["A. {1, 2, 3, 4, 5}", "B. {5, 4, 3, 2, 1}", "C. {1, 3, 2, 4, 5}", "D. {3, 1, 2, 4, 5}"], "ans": "A"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[6] = {2, 8, 8, 3, 8, 5};\nint k = 0;\nfor (int i = 1; i < 6; i++) {\n    if (a[i] >= a[k]) {\n        k = i;\n    }\n}\ncout << k;", "opts": ["A. 1", "B. 2", "C. 4", "D. 5"], "ans": "C"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[6] = {2, 8, 8, 3, 8, 5};\nint k = 0;\nfor (int i = 1; i < 6; i++) {\n    if (a[i] > a[k]) {\n        k = i;\n    }\n}\ncout << k;", "opts": ["A. 1", "B. 2", "C. 4", "D. 5"], "ans": "A"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[5] = {1, 3, 1, 3, 1};\nint ans = 0;\nfor (int i = 0; i < 5; i++) {\n    for (int j = i + 1; j < 5; j++) {\n        if (a[i] == a[j]) {\n            ans++;\n        }\n    }\n}\ncout << ans;", "opts": ["A. 3", "B. 4", "C. 5", "D. 6"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[5] = {4, 1, 3, 2, 5};\nint cnt = 0;\nfor (int i = 0; i < 5; i++) {\n    for (int j = i + 1; j < 5; j++) {\n        if (a[i] > a[j]) {\n            cnt++;\n        }\n    }\n}\ncout << cnt;", "opts": ["A. 2", "B. 3", "C. 4", "D. 5"], "ans": "C"},
    {"q": "下面程序片段执行后，数组 b 的内容是？", "code": "int a[5] = {1, 2, 3, 4, 5};\nint b[5];\nfor (int i = 0; i < 5; i++) {\n    b[i] = a[(i + 2) % 5];\n}", "opts": ["A. {1, 2, 3, 4, 5}", "B. {2, 3, 4, 5, 1}", "C. {3, 4, 5, 1, 2}", "D. {5, 4, 3, 2, 1}"], "ans": "C"},
    {"q": "下面程序片段执行后，数组 a 的内容是？", "code": "int a[6] = {1, 2, 3, 4, 5, 6};\nfor (int i = 0; i < 3; i++) {\n    int t = a[i];\n    a[i] = a[5 - i];\n    a[5 - i] = t;\n}", "opts": ["A. {6, 5, 4, 3, 2, 1}", "B. {4, 5, 6, 1, 2, 3}", "C. {1, 2, 3, 4, 5, 6}", "D. {6, 2, 3, 4, 5, 1}"], "ans": "A"},
    {"q": "下面二分查找程序的输出结果是？", "code": "int a[7] = {1, 4, 6, 9, 12, 15, 18};\nint l = 0, r = 6, x = 12, ans = -1;\nwhile (l <= r) {\n    int mid = (l + r) / 2;\n    if (a[mid] == x) {\n        ans = mid;\n        break;\n    } else if (a[mid] < x) {\n        l = mid + 1;\n    } else {\n        r = mid - 1;\n    }\n}\ncout << ans;", "opts": ["A. 3", "B. 4", "C. 5", "D. -1"], "ans": "B"},
    {"q": "下面二分查找程序中，第一次循环结束后 l 和 r 的值分别是？", "code": "int a[8] = {2, 4, 6, 8, 10, 12, 14, 16};\nint l = 0, r = 7, x = 5;\nint mid = (l + r) / 2;\nif (a[mid] < x) {\n    l = mid + 1;\n} else {\n    r = mid - 1;\n}", "opts": ["A. l = 0, r = 2", "B. l = 4, r = 7", "C. l = 0, r = 7", "D. l = 3, r = 7"], "ans": "A"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[3][3] = {\n    {3, 1, 2},\n    {6, 5, 4},\n    {7, 9, 8}\n};\nint s = 0;\nfor (int i = 0; i < 3; i++) {\n    s += a[i][i];\n}\ncout << s;", "opts": ["A. 13", "B. 16", "C. 18", "D. 21"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[3][3] = {\n    {3, 1, 2},\n    {6, 5, 4},\n    {7, 9, 8}\n};\nint s = 0;\nfor (int i = 0; i < 3; i++) {\n    s += a[i][2 - i];\n}\ncout << s;", "opts": ["A. 14", "B. 15", "C. 16", "D. 18"], "ans": "A"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[4] = {2, 4, 6, 8};\nfor (int i = 0; i < 4; i++) {\n    a[i] += i;\n}\ncout << a[1] << a[3];", "opts": ["A. 48", "B. 59", "C. 410", "D. 510"], "ans": "D"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[5] = {5, 1, 5, 2, 5};\nint sum = 0;\nfor (int i = 0; i < 5; i++) {\n    if (a[i] == a[0]) {\n        sum += i;\n    }\n}\ncout << sum;", "opts": ["A. 5", "B. 6", "C. 7", "D. 10"], "ans": "B"},
    {"q": "下面程序片段执行后，数组 a 的内容是？", "code": "int a[5] = {9, 7, 5, 3, 1};\nfor (int i = 0; i < 2; i++) {\n    int t = a[i];\n    a[i] = a[i + 3];\n    a[i + 3] = t;\n}", "opts": ["A. {3, 1, 5, 9, 7}", "B. {1, 3, 5, 7, 9}", "C. {9, 7, 5, 3, 1}", "D. {3, 7, 5, 9, 1}"], "ans": "A"},
    {"q": "下面程序片段的输出结果是？", "code": "int a[5] = {1, 4, 2, 5, 3};\nint mx = a[0] + a[1];\nfor (int i = 1; i < 4; i++) {\n    if (a[i] + a[i + 1] > mx) {\n        mx = a[i] + a[i + 1];\n    }\n}\ncout << mx;", "opts": ["A. 5", "B. 6", "C. 7", "D. 8"], "ans": "C"},
    {"q": "1 到 60 中能被 3 或 5 整除的整数有多少个？", "code": None, "opts": ["A. 28", "B. 30", "C. 32", "D. 36"], "ans": "A"},
    {"q": "1 到 80 中能被 4 或 6 整除的整数有多少个？", "code": None, "opts": ["A. 26", "B. 27", "C. 30", "D. 33"], "ans": "B"},
    {"q": "1 到 100 中既能被 2 整除又能被 5 整除的整数有多少个？", "code": None, "opts": ["A. 10", "B. 20", "C. 30", "D. 50"], "ans": "A"},
    {"q": "下面程序片段的输出结果是？", "code": "int cnt = 0;\nfor (int i = 1; i <= 30; i++) {\n    if (i % 2 != 0 && i % 3 != 0) {\n        cnt++;\n    }\n}\ncout << cnt;", "opts": ["A. 8", "B. 10", "C. 12", "D. 15"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int cnt = 0;\nfor (int i = 1; i <= 50; i++) {\n    if (i % 4 == 0 || i % 7 == 0) {\n        cnt++;\n    }\n}\ncout << cnt;", "opts": ["A. 16", "B. 17", "C. 18", "D. 19"], "ans": "B"},
    {"q": "1 到 100 中不能被 3 整除的整数有多少个？", "code": None, "opts": ["A. 33", "B. 34", "C. 66", "D. 67"], "ans": "D"},
    {"q": "1 到 120 中能被 6 整除但不能被 4 整除的整数有多少个？", "code": None, "opts": ["A. 10", "B. 15", "C. 20", "D. 30"], "ans": "A"},
    {"q": "下面程序片段的输出结果是？", "code": "int ans = 0;\nfor (int i = 1; i <= 40; i++) {\n    if (i % 2 == 0 && i % 5 != 0) {\n        ans++;\n    }\n}\ncout << ans;", "opts": ["A. 12", "B. 14", "C. 16", "D. 20"], "ans": "C"},
    {"q": "1 到 90 中能被 2、3、5 中至少一个整除的整数有多少个？", "code": None, "opts": ["A. 60", "B. 66", "C. 68", "D. 72"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int cnt = 0;\nfor (int i = 10; i <= 99; i++) {\n    if (i % 10 == 3 || i / 10 == 3) {\n        cnt++;\n    }\n}\ncout << cnt;", "opts": ["A. 10", "B. 18", "C. 19", "D. 20"], "ans": "B"},
    {"q": "1 到 50 中是 3 的倍数但不是 2 的倍数的整数有多少个？", "code": None, "opts": ["A. 8", "B. 9", "C. 16", "D. 25"], "ans": "A"},
    {"q": "下面程序片段的输出结果是？", "code": "int cnt = 0;\nfor (int i = 1; i <= 36; i++) {\n    if (36 % i == 0) {\n        cnt++;\n    }\n}\ncout << cnt;", "opts": ["A. 6", "B. 8", "C. 9", "D. 12"], "ans": "C"},
    {"q": "1 到 72 中能被 8 或 9 整除的整数有多少个？", "code": None, "opts": ["A. 15", "B. 16", "C. 17", "D. 18"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 9;\ncout << (x % 3 == 0 && x % 2 == 0 || x > 8);", "opts": ["A. 0", "B. 1", "C. 9", "D. 编译错误"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 10;\ncout << (x > 5 || x / 0 == 1);", "opts": ["A. 0", "B. 1", "C. 运行时错误", "D. 编译错误"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 3;\ncout << (x < 0 && x / 0 == 1);", "opts": ["A. 0", "B. 1", "C. 运行时错误", "D. 编译错误"], "ans": "A"},
    {"q": "下面程序片段的输出结果是？", "code": "int a = 4, b = 7;\ncout << (a + b > 10 && a * b < 30);", "opts": ["A. 0", "B. 1", "C. 11", "D. 28"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 12;\ncout << (!(x % 3 == 0) || x % 4 == 0);", "opts": ["A. 0", "B. 1", "C. 3", "D. 4"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 15;\ncout << (x % 3 == 0 && !(x % 5 == 0));", "opts": ["A. 0", "B. 1", "C. 3", "D. 5"], "ans": "A"},
    {"q": "下面程序片段的输出结果是？", "code": "int a = 2, b = 3, c = 4;\ncout << (a + b * c == 14 || (a + b) * c == 20);", "opts": ["A. 0", "B. 1", "C. 14", "D. 20"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 8;\nbool ok = x > 5 && x < 10 && x % 2 == 0;\ncout << ok;", "opts": ["A. 0", "B. 1", "C. true", "D. false"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 6;\nbool ok = x % 2 == 0 || x % 3 == 0 && x % 5 == 0;\ncout << ok;", "opts": ["A. 0", "B. 1", "C. true", "D. false"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 7;\nbool ok = (x % 2 == 0 || x % 3 == 1) && x < 10;\ncout << ok;", "opts": ["A. 0", "B. 1", "C. 7", "D. 10"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 5, y = 2;\ncout << (x / y == 2 && x % y == 0);", "opts": ["A. 0", "B. 1", "C. 2", "D. 5"], "ans": "A"},
    {"q": "下面程序片段的输出结果是？", "code": "int a = 1, b = 0;\ncout << (a || b && !a);", "opts": ["A. 0", "B. 1", "C. true", "D. false"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int a = 0, b = 1;\ncout << (!a && b || a);", "opts": ["A. 0", "B. 1", "C. true", "D. false"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 20;\ncout << (x % 4 == 0 && x % 6 == 0 || x % 10 == 0);", "opts": ["A. 0", "B. 1", "C. 4", "D. 10"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 11;\ncout << (x > 10 ? x % 3 == 2 : x % 2 == 0);", "opts": ["A. 0", "B. 1", "C. 2", "D. 11"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 4;\ncout << ((x > 3 ? x + 1 : x - 1) == 5);", "opts": ["A. 0", "B. 1", "C. 4", "D. 5"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int a = 3, b = 4;\ncout << (a * a + b * b == 25 && a + b == 7);", "opts": ["A. 0", "B. 1", "C. 7", "D. 25"], "ans": "B"},
    {"q": "下面程序片段的输出结果是？", "code": "int x = 14;\ncout << (!(x % 7 == 0 && x % 2 == 0));", "opts": ["A. 0", "B. 1", "C. 2", "D. 7"], "ans": "A"},
]


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("小学组提高级 C++ 排序、整除统计与逻辑专项训练（50题）")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(16)
    run.bold = True

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("针对错题 9、11、25、37：排序过程、数组状态、整除计数、逻辑表达式优先级")
    nr.font.name = "Microsoft YaHei"
    nr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(89, 99, 110)

    for i, q in enumerate(questions, 1):
        add_question(doc, i, q)

    doc.add_page_break()
    ans_title = doc.add_paragraph()
    ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = ans_title.add_run("答案")
    ar.font.name = "Microsoft YaHei"
    ar._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    ar.font.size = Pt(16)
    ar.bold = True

    table = doc.add_table(rows=11, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["1-10", "11-20", "21-30", "31-40", "41-50"]
    for c, h in enumerate(headers):
        cell = table.cell(0, c)
        set_cell_shading(cell, "DDEBFF")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10)

    for row in range(1, 11):
        for col in range(5):
            num = row + col * 10
            cell = table.cell(row, col)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"{num}. {questions[num - 1]['ans']}")
            r.font.name = "Microsoft YaHei"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            r.font.size = Pt(10.5)

    doc.save(OUT)
    print(OUT)
    print(f"questions={len(questions)} code_questions={sum(1 for q in questions if q.get('code'))}")


if __name__ == "__main__":
    build_doc()
