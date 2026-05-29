from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from make_remedial_50_cpp_doc import add_question, add_code_line, set_cell_shading
from make_gesp_style_paper import add_heading, add_label_para, set_font


OUT = "/Users/cen/Documents/Codex/2026-05-18/4-5-1-2-3-4/GESP四级偏上五级入门模拟卷2_10选4编程_样例已校验.docx"


mc_questions = [
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int a = 19, b = 4;\ncout << a / b + a % b * 2;",
        "opts": ["A. 7", "B. 10", "C. 11", "D. 19"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int cnt = 0;\nfor (int i = 1; i <= 5; i++) {\n    for (int j = 1; j <= i; j++) {\n        if ((i + j) % 2 == 0) {\n            cnt++;\n        }\n    }\n}\ncout << cnt;",
        "opts": ["A. 6", "B. 8", "C. 9", "D. 15"],
        "ans": "C",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "string s = \"p10q25\";\nint x = 0;\nfor (int i = 0; i < s.size(); i++) {\n    if (s[i] >= '0' && s[i] <= '9') {\n        x = x * 10 + (s[i] - '0');\n    }\n}\ncout << x;",
        "opts": ["A. 35", "B. 1025", "C. 125", "D. 10025"],
        "ans": "B",
    },
    {
        "q": "用筛法标记 2 到 70 的合数，最后未被标记的数有多少个？",
        "code": None,
        "opts": ["A. 18", "B. 19", "C. 20", "D. 21"],
        "ans": "B",
    },
    {
        "q": "执行下面一趟冒泡排序后，数组 a 的内容是？",
        "code": "int a[6] = {6, 1, 5, 2, 4, 3};\nfor (int j = 0; j < 5; j++) {\n    if (a[j] > a[j + 1]) {\n        int t = a[j];\n        a[j] = a[j + 1];\n        a[j + 1] = t;\n    }\n}",
        "opts": ["A. {1, 5, 2, 4, 3, 6}", "B. {1, 2, 3, 4, 5, 6}", "C. {1, 6, 2, 4, 3, 5}", "D. {6, 5, 4, 3, 2, 1}"],
        "ans": "A",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int a[3][3] = {\n    {5, 1, 2},\n    {4, 7, 3},\n    {6, 8, 9}\n};\nint s = 0;\nfor (int i = 0; i < 3; i++) {\n    s += a[i][2 - i];\n}\ncout << s;",
        "opts": ["A. 12", "B. 15", "C. 18", "D. 21"],
        "ans": "B",
    },
    {
        "q": "下面函数调用 f(6) 的返回值是？",
        "code": "int f(int n) {\n    if (n <= 2) {\n        return 1;\n    }\n    return f(n - 1) + f(n - 2);\n}",
        "opts": ["A. 5", "B. 8", "C. 13", "D. 21"],
        "ans": "B",
    },
    {
        "q": "下面二分查找程序的输出结果是？",
        "code": "int a[7] = {2, 4, 6, 8, 10, 12, 14};\nint l = 0, r = 6, x = 10, ans = -1;\nwhile (l <= r) {\n    int mid = (l + r) / 2;\n    if (a[mid] == x) {\n        ans = mid;\n        break;\n    } else if (a[mid] < x) {\n        l = mid + 1;\n    } else {\n        r = mid - 1;\n    }\n}\ncout << ans;",
        "opts": ["A. 3", "B. 4", "C. 5", "D. -1"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int x = 18;\ncout << (x % 2 == 0 && x % 5 == 0 || x % 3 == 0);",
        "opts": ["A. 0", "B. 1", "C. 2", "D. 3"],
        "ans": "B",
    },
    {
        "q": "从数字 1、2、3、4、5 中选出两个不同数字组成两位数，其中偶数有多少个？",
        "code": None,
        "opts": ["A. 6", "B. 8", "C. 10", "D. 20"],
        "ans": "B",
    },
]


def solve_row_col(inp: str) -> str:
    nums = list(map(int, inp.split()))
    n, m = nums[0], nums[1]
    idx = 2
    a = []
    for _ in range(n):
        a.append(nums[idx:idx + m])
        idx += m
    x, y = nums[idx] - 1, nums[idx + 1] - 1
    ans = sum(a[x]) + sum(a[i][y] for i in range(n)) - a[x][y]
    return str(ans)


def solve_prime_count(inp: str) -> str:
    l, r = map(int, inp.split())
    if r < 2:
        return "0"
    is_prime = [True] * (r + 1)
    is_prime[0] = False
    if r >= 1:
        is_prime[1] = False
    p = 2
    while p * p <= r:
        if is_prime[p]:
            for j in range(p * p, r + 1, p):
                is_prime[j] = False
        p += 1
    return str(sum(1 for x in range(l, r + 1) if is_prime[x]))


def solve_tasks(inp: str) -> str:
    nums = list(map(int, inp.split()))
    n, t = nums[0], nums[1]
    arr = sorted(nums[2:2 + n])
    used = 0
    total = 0
    for x in arr:
        if total + x <= t:
            total += x
            used += 1
        else:
            break
    return str(used)


def solve_longest_rise(inp: str) -> str:
    nums = list(map(int, inp.split()))
    n = nums[0]
    a = nums[1:1 + n]
    best = cur = 1
    for i in range(1, n):
        if a[i] > a[i - 1]:
            cur += 1
        else:
            cur = 1
        best = max(best, cur)
    return str(best)


programs = [
    {
        "title": "矩阵十字和",
        "background": "给定一个 n 行 m 列的整数矩阵，以及一个位置 (x, y)。请输出第 x 行与第 y 列所有元素之和，交叉位置只计算一次。行列编号从 1 开始。",
        "input": "第一行输入 n 和 m。接下来 n 行，每行 m 个整数。最后一行输入 x 和 y。",
        "output": "输出一个整数，表示矩阵十字和。",
        "constraints": "1 <= n, m <= 50；矩阵元素为 0 到 1000 的整数。",
        "solver": solve_row_col,
        "inputs": [
            "3 3\n1 2 3\n4 5 6\n7 8 9\n2 2",
            "2 4\n1 1 1 1\n2 2 2 2\n1 3",
            "4 2\n3 5\n7 9\n1 4\n6 8\n4 1",
            "1 5\n2 4 6 8 10\n1 4",
            "2 2\n10 20\n30 40\n2 2",
        ],
    },
    {
        "title": "区间质数个数",
        "background": "给定两个整数 L 和 R，请统计闭区间 [L, R] 中质数的个数。",
        "input": "输入一行，包含两个整数 L 和 R。",
        "output": "输出一个整数，表示区间内质数的个数。",
        "constraints": "1 <= L <= R <= 100000。",
        "solver": solve_prime_count,
        "inputs": [
            "1 10",
            "14 20",
            "30 50",
            "90 100",
            "997 1009",
        ],
    },
    {
        "title": "最多完成任务数",
        "background": "有 n 个任务，每个任务需要一定时间。小杨总时间为 T，每个任务最多完成一次。为了完成尽量多的任务，请输出最多能完成几个任务。",
        "input": "第一行输入 n 和 T。第二行输入 n 个正整数，表示每个任务需要的时间。",
        "output": "输出一个整数，表示最多能完成的任务数量。",
        "constraints": "1 <= n <= 1000；1 <= T <= 100000；每个任务耗时为 1 到 100000。",
        "solver": solve_tasks,
        "inputs": [
            "5 10\n1 3 4 6 8",
            "6 12\n5 1 2 7 3 4",
            "5 7\n8 9 10 11 12",
            "7 15\n1 1 2 2 3 8 9",
            "8 9\n4 4 4 1 1 1 10 2",
        ],
    },
    {
        "title": "最长连续上升段",
        "background": "给定一个长度为 n 的整数序列，请求出最长连续上升段的长度。连续上升段要求相邻后一个数严格大于前一个数。",
        "input": "第一行输入 n。第二行输入 n 个整数。",
        "output": "输出一个整数，表示最长连续上升段的长度。",
        "constraints": "1 <= n <= 100000；序列元素为 int 范围内整数。",
        "solver": solve_longest_rise,
        "inputs": [
            "5\n1 2 3 2 5",
            "6\n6 5 4 3 2 1",
            "7\n1 3 5 7 9 11 13",
            "8\n2 2 3 4 1 2 3 0",
            "10\n1 2 1 2 3 4 0 1 2 3",
        ],
    },
]


def attach_verified_tests():
    for item in programs:
        tests = []
        for inp in item["inputs"]:
            out = item["solver"](inp)
            tests.append((inp, out))
        item["tests"] = tests


def add_label_para(doc, label, body):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.25)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label)
    set_font(r1, 10.5, True)
    r2 = p.add_run(body)
    set_font(r2, 10.5)


def add_tests_table(doc, tests):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["序号", "输入", "输出"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "DDEBFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, 9.5, True)
    for idx, (inp, out) in enumerate(tests, 1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = inp
        cells[2].text = out
        for ci, c in enumerate(cells):
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = "Courier New" if ci in (1, 2) else "Microsoft YaHei"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), r.font.name)
                    r.font.size = Pt(8.5)
    doc.add_paragraph()


def build_doc():
    attach_verified_tests()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("GESP C++ 四级偏上 / 五级入门 模拟卷（二）")
    set_font(r, 16, True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("10 道选择题 + 4 道编程题；每道编程题 5 组测试例，测试输出已由脚本校验")
    set_font(nr, 9, False, "59636E")

    add_heading(doc, "一、选择题（每题 3 分，共 30 分）")
    for i, q in enumerate(mc_questions, 1):
        add_question(doc, i, q)

    add_heading(doc, "二、编程题（每题 25 分，共 100 分）")
    for idx, item in enumerate(programs, 1):
        add_heading(doc, f"编程题 {idx}. {item['title']}", 12)
        add_label_para(doc, "题目描述：", item["background"])
        add_label_para(doc, "输入格式：", item["input"])
        add_label_para(doc, "输出格式：", item["output"])
        add_label_para(doc, "数据范围：", item["constraints"])
        add_label_para(doc, "测试例：", "")
        add_tests_table(doc, item["tests"])

    doc.add_page_break()
    add_heading(doc, "答案")
    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(10):
        row = i // 5
        col = i % 5
        cell = table.cell(row, col)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{i + 1}. {mc_questions[i]['ans']}")
        set_font(r, 10.5, True)

    p = doc.add_paragraph()
    r = p.add_run("编程题测试例输出已按题意用脚本重新计算校验；正式考试建议另加隐藏测试。")
    set_font(r, 10.5)

    doc.save(OUT)
    print(OUT)
    print(f"mc={len(mc_questions)} programs={len(programs)} tests={sum(len(p['tests']) for p in programs)}")


if __name__ == "__main__":
    build_doc()
