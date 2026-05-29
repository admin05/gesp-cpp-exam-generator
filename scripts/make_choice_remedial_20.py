from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from make_remedial_50_cpp_doc import add_question, set_cell_shading
from make_gesp_style_paper import set_font, add_heading


OUT = "/Users/cen/Documents/Codex/2026-05-18/4-5-1-2-3-4/GESP四级偏上五级入门_选择题错题专项20题.docx"


questions = [
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int a = 8, b = 3, c = 2;\ncout << a - b * c + a / c;",
        "opts": ["A. 6", "B. 8", "C. 10", "D. 14"],
        "ans": "A",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int x = 17;\ncout << x / 4 + x % 4 * 2;",
        "opts": ["A. 5", "B. 6", "C. 8", "D. 10"],
        "ans": "B",
    },
    {
        "q": "下面表达式的值是？",
        "code": "18 / 5 + 18 % 5 + 2 * 3",
        "opts": ["A. 9", "B. 10", "C. 12", "D. 15"],
        "ans": "C",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int cnt = 0;\nfor (int i = 1; i <= 5; i++) {\n    for (int j = i; j <= 5; j++) {\n        if ((i + j) % 2 == 1) {\n            cnt++;\n        }\n    }\n}\ncout << cnt;",
        "opts": ["A. 4", "B. 6", "C. 8", "D. 10"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int ans = 0;\nfor (int i = 1; i <= 4; i++) {\n    for (int j = 1; j <= 4; j++) {\n        if (i < j && (i + j) % 2 == 0) {\n            ans++;\n        }\n    }\n}\ncout << ans;",
        "opts": ["A. 2", "B. 3", "C. 4", "D. 6"],
        "ans": "A",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int cnt = 0;\nfor (int i = 1; i <= 3; i++) {\n    for (int j = 1; j <= 4; j++) {\n        if (i * j % 2 == 0) {\n            cnt++;\n        }\n    }\n}\ncout << cnt;",
        "opts": ["A. 6", "B. 8", "C. 9", "D. 12"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "string s = \"x20y31\";\nint sum = 0;\nfor (int i = 0; i < s.size(); i++) {\n    if (s[i] >= '0' && s[i] <= '9') {\n        sum += s[i] - '0';\n    }\n}\ncout << sum;",
        "opts": ["A. 6", "B. 20", "C. 31", "D. 2031"],
        "ans": "A",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "string s = \"a05b7\";\nint x = 0;\nfor (int i = 0; i < s.size(); i++) {\n    if (s[i] >= '0' && s[i] <= '9') {\n        x = x * 10 + (s[i] - '0');\n    }\n}\ncout << x;",
        "opts": ["A. 12", "B. 57", "C. 507", "D. 057"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "string s = \"314159\";\nint cnt = 0;\nfor (int i = 0; i < s.size(); i++) {\n    if ((s[i] - '0') % 2 == 1) {\n        cnt++;\n    }\n}\ncout << cnt;",
        "opts": ["A. 3", "B. 4", "C. 5", "D. 6"],
        "ans": "C",
    },
    {
        "q": "用筛法标记 2 到 60 的合数，最后未被标记的数有多少个？",
        "code": None,
        "opts": ["A. 16", "B. 17", "C. 18", "D. 19"],
        "ans": "B",
    },
    {
        "q": "下面筛法程序输出结果是？",
        "code": "bool vis[31] = {};\nfor (int i = 2; i <= 30; i++) {\n    if (!vis[i]) {\n        for (int j = i * i; j <= 30; j += i) {\n            vis[j] = true;\n        }\n    }\n}\nint cnt = 0;\nfor (int i = 2; i <= 30; i++) {\n    if (!vis[i]) {\n        cnt++;\n    }\n}\ncout << cnt;",
        "opts": ["A. 8", "B. 9", "C. 10", "D. 11"],
        "ans": "C",
    },
    {
        "q": "下面程序片段判断 n 是否为质数。若 n = 91，第一次使 ok 变为 false 的 i 是？",
        "code": "int n = 91;\nbool ok = true;\nfor (int i = 2; i * i <= n; i++) {\n    if (n % i == 0) {\n        ok = false;\n        break;\n    }\n}",
        "opts": ["A. 7", "B. 9", "C. 13", "D. 91"],
        "ans": "A",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int a[5] = {3, 1, 4, 1, 5};\nint mx = a[0];\nfor (int i = 1; i < 5; i++) {\n    if (a[i] > mx) {\n        mx = a[i];\n    }\n}\ncout << mx;",
        "opts": ["A. 1", "B. 3", "C. 4", "D. 5"],
        "ans": "D",
    },
    {
        "q": "执行下面一趟冒泡排序后，数组 a 的内容是？",
        "code": "int a[5] = {5, 2, 4, 1, 3};\nfor (int j = 0; j < 4; j++) {\n    if (a[j] > a[j + 1]) {\n        int t = a[j];\n        a[j] = a[j + 1];\n        a[j + 1] = t;\n    }\n}",
        "opts": ["A. {2, 4, 1, 3, 5}", "B. {1, 2, 3, 4, 5}", "C. {2, 5, 1, 3, 4}", "D. {5, 4, 3, 2, 1}"],
        "ans": "A",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int a[3][3] = {\n    {1, 2, 3},\n    {4, 5, 6},\n    {7, 8, 9}\n};\nint s = 0;\nfor (int i = 0; i < 3; i++) {\n    s += a[i][2 - i];\n}\ncout << s;",
        "opts": ["A. 12", "B. 15", "C. 18", "D. 21"],
        "ans": "B",
    },
    {
        "q": "下面函数调用 f(4) 的返回值是？",
        "code": "int f(int n) {\n    if (n == 1) {\n        return 1;\n    }\n    return n * f(n - 1);\n}",
        "opts": ["A. 4", "B. 10", "C. 16", "D. 24"],
        "ans": "D",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int x = 12;\ncout << (x % 3 == 0 && x % 5 == 0 || x % 4 == 0);",
        "opts": ["A. 0", "B. 1", "C. 3", "D. 4"],
        "ans": "B",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int x = 7;\ncout << (x > 5 ? x % 3 : x % 2);",
        "opts": ["A. 1", "B. 2", "C. 3", "D. 7"],
        "ans": "A",
    },
    {
        "q": "下面二分查找程序的输出结果是？",
        "code": "int a[6] = {1, 3, 5, 7, 9, 11};\nint l = 0, r = 5, x = 8, ans = -1;\nwhile (l <= r) {\n    int mid = (l + r) / 2;\n    if (a[mid] == x) {\n        ans = mid;\n        break;\n    } else if (a[mid] < x) {\n        l = mid + 1;\n    } else {\n        r = mid - 1;\n    }\n}\ncout << ans;",
        "opts": ["A. -1", "B. 2", "C. 3", "D. 4"],
        "ans": "A",
    },
    {
        "q": "下面程序片段的输出结果是？",
        "code": "int cnt = 0;\nfor (int i = 1; i <= 24; i++) {\n    if (i % 2 == 0 || i % 3 == 0) {\n        cnt++;\n    }\n}\ncout << cnt;",
        "opts": ["A. 12", "B. 16", "C. 18", "D. 20"],
        "ans": "B",
    },
]


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("GESP C++ 四级偏上 / 五级入门 选择题错题专项（20题）")
    set_font(r, 16, True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("重点回炉：运算优先级、嵌套循环计数、字符串数字处理、筛法与质数判断")
    set_font(nr, 9, False, "59636E")

    add_heading(doc, "选择题（每题 3 分，共 60 分）")
    for i, q in enumerate(questions, 1):
        add_question(doc, i, q)

    doc.add_page_break()
    add_heading(doc, "答案")
    table = doc.add_table(rows=4, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(20):
        row = i // 5
        col = i % 5
        cell = table.cell(row, col)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if row == 0:
            set_cell_shading(cell, "F6F8FA")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{i + 1}. {questions[i]['ans']}")
        set_font(r, 10.5, True)

    doc.save(OUT)
    print(OUT)
    print(f"questions={len(questions)} remedial={12}")


if __name__ == "__main__":
    build_doc()
